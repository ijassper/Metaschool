# 엑셀/워드 출력 (submission_export_excel 등)

from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.utils import timezone
from urllib.parse import quote

# 엑셀 및 워드 라이브러리
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 커스텀 데코레이터 및 모델 임포트
from accounts.decorators import teacher_required
from ..models import Activity, Answer

# [1] 학생 순수 답안 엑셀 내보내기 (제출 현황 페이지용)
@login_required
@teacher_required
def submission_export_excel(request, activity_id):
    activity = get_object_or_404(Activity, id=activity_id, teacher=request.user)
    question = activity.questions.first() # 활동의 문항 정보
    
    # 1. 엑셀 파일 생성
    wb = Workbook()
    ws = wb.active
    ws.title = "학생답안목록"
    
    # 2. 헤더 작성 (요청하신 순서: 학년, 반, 번호, 이름, 답안)
    headers = ['학년', '반', '번호', '이름', '학생 답안']
    ws.append(headers)
    
    # 3. 데이터 추출 (해당 활동의 대상 학생 전체)
    students = activity.target_students.all().order_by('grade', 'class_no', 'number')
    
    for s in students:
        answer = Answer.objects.filter(student=s, question=question).first()
        
        # 학생이 작성한 답안이 있으면 가져오고, 없으면 공백처리
        # (통합 필드인 content를 사용하여 항목 1,2,3 합본을 가져옵니다)
        content = answer.content if answer and answer.content else ""
        
        ws.append([s.grade, s.class_no, s.number, s.name, content])
    
    # 4. 파일 다운로드 응답 생성
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    filename = f"학생답안_{activity.title}_{timezone.now().strftime('%m%d')}.xlsx"
    
    # 한글 파일명 깨짐 방지
    from urllib.parse import quote
    response['Content-Disposition'] = f"attachment; filename*=UTF-8''{quote(filename)}"
    wb.save(response)
    
    return response

# [2] AI 분석 결과 통합 엑셀 내보내기 (결과 분석 페이지용)
@login_required
@teacher_required
def analysis_export_excel(request, activity_id):
    activity = get_object_or_404(Activity, id=activity_id, teacher=request.user)
    question = activity.questions.first()
    
    # 1. 엑셀 파일 생성
    wb = Workbook()
    ws = wb.active
    ws.title = "AI 분석 결과"
    
    # 1. 헤더 동적 생성 (설계도 반영)
    headers = ['학년', '반', '번호', '이름']
    headers.append(activity.q1_title)
    if activity.q2_title: headers.append(activity.q2_title)
    if activity.q3_title: headers.append(activity.q3_title)
    headers.extend(['AI 분석 결과', '분석일시'])
    
    ws.append(headers)
    
    # 2. 데이터 추출
    students = activity.target_students.all().order_by('grade', 'class_no', 'number')
    for s in students:
        answer = Answer.objects.filter(student=s, question=question).first()
        
        row = [s.grade, s.class_no, s.number, s.name]
        
        if answer:
            row.append(answer.ans_q1 if answer.ans_q1 else "")
            if activity.q2_title: row.append(answer.ans_q2 if answer.ans_q2 else "")
            if activity.q3_title: row.append(answer.ans_q3 if answer.ans_q3 else "")
            row.append(answer.ai_result if answer.ai_result else "")
            row.append(answer.ai_updated_at.strftime('%Y-%m-%d %H:%M') if answer.ai_updated_at else "")
        else:
            # 미제출 학생 처리
            row.append("(미제출)")
            if activity.q2_title: row.append("")
            if activity.q3_title: row.append("")
            row.extend(["", ""])
            
        ws.append(row)
    
    # 4. 파일 다운로드 응답 생성
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    filename = f"AI_분석결과_{activity.title}_{timezone.now().strftime('%m%d')}.xlsx"
    
    # 한글 파일명 깨짐 방지 처리
    response['Content-Disposition'] = f"attachment; filename*=UTF-8''{quote(filename)}"
    wb.save(response)
    
    return response

# [3] 개별 답안지 워드 출력 (인쇄용 양식)
@login_required
@teacher_required
def export_answer_sheets_docx(request, activity_id):
    activity = get_object_or_404(Activity, id=activity_id, teacher=request.user)
    
    # 1. 워드 문서 생성 및 기본 설정
    doc = Document()
    
    # 2. 모든 응시 대상 학생을 순회
    students = activity.target_students.all().order_by('grade', 'class_no', 'number')
    
    for idx, student in enumerate(students):
        if idx > 0:
            doc.add_page_break() # 학생별로 페이지 나누기 (인쇄 편의성)

        # [표 생성] 4행 4열 (스크린샷 양식 반영)
        table = doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'
        
        # 표의 각 셀 너비 조정 (선택 사항)
        # table.autofit = False
        
        # 1행: 응시 과목 / 평가명
        table.cell(0, 0).text = "응시 과목"
        table.cell(0, 1).text = activity.section # 과목명
        table.cell(0, 2).text = "평가명"
        table.cell(0, 3).text = activity.title # 주제

        # 2행: 응시자 정보 / 응시 일시
        answer = activity.get_student_answer(student)
        submitted_at = answer.submitted_at.strftime('%Y-%m-%d %H:%M') if answer and answer.submitted_at else "-"
        
        table.cell(1, 0).text = "응시자 정보"
        table.cell(1, 1).text = f"{student.grade}학년 {student.class_no}반 {student.number}번 {student.name}"
        table.cell(1, 2).text = "응시 일시"
        table.cell(1, 3).text = submitted_at

        # 3행: 제목 (병합 처리)
        title_cell = table.cell(2, 0)
        title_cell.merge(table.cell(2, 3))
        title_cell.text = f"학생 답안 (제목: {activity.q1_title})"
        title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 4행: 답안 내용 (병합 처리 및 내용 삽입)
        content_cell = table.cell(3, 0)
        content_cell.merge(table.cell(3, 3))
        
        if answer:
            # 항목 1, 2, 3이 있다면 합쳐서 출력
            combined_content = ""
            if answer.ans_q1: combined_content += f"[{activity.q1_title}]\n{answer.ans_q1}\n\n"
            if answer.ans_q2: combined_content += f"[{activity.q2_title}]\n{answer.ans_q2}\n\n"
            if answer.ans_q3: combined_content += f"[{activity.q3_title}]\n{answer.ans_q3}"
            
            content_cell.text = combined_content if combined_content else answer.content
        else:
            content_cell.text = "\n\n(미제출된 답안입니다.)\n\n"

        # [디테일] 표 내부 텍스트 폰트 설정 (한글 깨짐 방지)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        
                        # --- [안전한 한글 폰트 설정 로직] ---
                        r = run._element
                        rPr = r.get_or_add_rPr() # 속성 요소가 없으면 생성
                        rFonts = rPr.get_or_add_rFonts() # 폰트 요소가 없으면 생성
                        rFonts.set(qn('w:eastAsia'), '함초롬바탕') # 한글 폰트 강제 지정
                        # -----------------------------------

    # 3. 파일 다운로드 응답 생성
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"답안지_목록_{activity.title}.docx"
    response['Content-Disposition'] = f"attachment; filename*=UTF-8''{quote(filename)}"
    doc.save(response)
    
    return response

@login_required
@teacher_required
def print_answer_sheets(request, activity_id):
    activity = get_object_or_404(Activity, id=activity_id, teacher=request.user)
    students = activity.target_students.all().order_by('grade', 'class_no', 'number')
    
    # 각 학생별 답안 데이터를 미리 매핑해서 전달
    for s in students:
        s.my_answer = activity.get_student_answer(s)
        
    return render(request, 'activities/print_answers.html', {
        'activity': activity,
        'students': students,
    })